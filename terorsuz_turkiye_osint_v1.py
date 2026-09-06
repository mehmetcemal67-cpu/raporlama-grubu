import os
import re
import json
import logging
from datetime import datetime
from typing import List, Dict, Any, Optional

import networkx as nx
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

# Logging Yapılandırması
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger("TerorsuzTurkiyePipeline")

# ==============================================================================
# 1. BÖLÜM: KAYNAK VE RADAR VERİ HAVUZU TANIMLARI
# ==============================================================================

# Kürt Bölgesel / Odaklı Medya Havuzu (Ağırlığı artırılmış kaynaklar)
KURDISH_MEDIA_SOURCES = [
    "Rudaw", "Darka Mazi", "Shafaq News", "Kurdistan24", "Kurdpress", 
    "The New Region", "Ilke TV", "Yeni Yaşam", "Sterk TV", "Rojnews", "Amargi"
]

# Yabancı / Uluslararası Basın Havuzu
INTERNATIONAL_MEDIA_SOURCES = [
    "Jerusalem Post", "Reuters", "AP News", "Al-Monitor", "Middle East Eye", 
    "BBC World", "Financial Times", "The Hindu", "Telesur English"
]

# Yerli Ana Akım ve Sosyal Medya Kanalları
MAINSTREAM_SOURCES = [
    "Manşet Haber", "Büyük Sivas Haber", "soL Haber", "Medyascope", 
    "Bengü Türk", "Pusu Haber", "Odatv", "Yeni Şafak", "Yeniçağ Gazetesi"
]

# Sıkı Alan Anahtar Kelime Whitelist (Güvenlik / Terör / Bölgesel Siyaset)
SECURITY_DOMAIN_KEYWORDS = [
    "öcalan", "pkk", "kck", "sdg", "ypg", "pyd", "dem parti", "terörsüz türkiye",
    "imralı", "kandil", "suriye", "barzani", "kdp", "mhp", "chp", "sürec", "çözüm",
    "silahsızlanma", "fesih", "kurdistan", "mazlum abdi", "cemil bayık", "mustafa karasu",
    "shafaq", "darka mazi", "rudaw", "rojava", "entegrasyon", "anadilde eğitim", 
    "milli dayanışma", "7595", "türkmen", "hbdh", "mlkp", "yıldıray oğur", "jonathan spyer"
]

# Kesinlikle Elenecek Magazin / Trend / Alakasız Kalıplar (Blacklist)
IRRELEVANT_BLACKLIST = [
    "scary movie", "trailer", "official music video", "gameplay", "full movie",
    "fragman", "komedi", "film izle", "parodi", "tiktok dance", "magazin", 
    "futbol", "süper lig", "maç özeti"
]


# ==============================================================================
# 2. BÖLÜM: SÜZGEÇ VE FİLTRELEME MİMARİSİ ("Scary Movie" Vb. İzolasyonu)
# ==============================================================================

class ContentFilter:
    """
    'Şu An Bilmem Gerekenler' bölümüne alakasız trend veya magazin 
    içeriklerinin girmesini engelleyen filtreleme sınıfı.
    """
    @staticmethod
    def is_relevant(title: str, summary: str = "", content: str = "") -> bool:
        full_text = f"{title} {summary} {content}".lower()
        
        # 1. Blacklist kontrolü (Scary Movie vb.)
        if any(bad_word in full_text for bad_word in IRRELEVANT_BLACKLIST):
            return False
            
        # 2. Domain Whitelist kontrolü
        return any(keyword in full_text for keyword in SECURITY_DOMAIN_KEYWORDS)

    @classmethod
    def filter_delta_scan(cls, raw_new_items: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """
        Son iki tarama arasındaki farktan (delta) gelen içerikleri süzer.
        """
        filtered = []
        for item in raw_new_items:
            t = item.get("title", "")
            s = item.get("summary", "")
            c = item.get("content", "")
            if cls.is_relevant(t, s, c):
                filtered.append(item)
            else:
                logger.info(f"Filtreye takılan alakasız içerik izole edildi: {t}")
        return filtered


# ==============================================================================
# 3. BÖLÜM: GÜNDEM VE SÖYLEM HARİTASI (Kürt Medyası & Karşıt Söylem Odaklı)
# ==============================================================================

class DiscourseClusteringEngine:
    """
    Kürt medyasını ve karşıt söylemleri odağa alan gündem kümeleme motoru.
    """
    @staticmethod
    def calculate_cross_perspective_score(cluster_items: List[Dict[str, Any]]) -> float:
        sources = set(item.get("source_name", "") for item in cluster_items)
        
        has_kurdish = any(s in KURDISH_MEDIA_SOURCES for s in sources)
        has_international = any(s in INTERNATIONAL_MEDIA_SOURCES for s in sources)
        has_mainstream = any(s in MAINSTREAM_SOURCES for s in sources)
        
        base_score = len(cluster_items)
        
        # Karşıt söylem ve Kürt medyası barındıran gündemlerin öncelik ağırlığını artır
        if has_kurdish and (has_mainstream or has_international):
            return base_score * 3.5  # Yüksek karşıtlık bonusu
        elif has_kurdish:
            return base_score * 2.2  # Odak Kürt medyası bonusu
        
        return float(base_score)

    @classmethod
    def build_agenda_map(cls, scraped_feed: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        # Filtreleme
        clean_feed = [item for item in scraped_feed if ContentFilter.is_relevant(
            item.get("title", ""), item.get("summary", ""), item.get("content", "")
        )]
        
        # Kümeleme (Cluster) Eşleştirmesi
        clusters: Dict[str, List[Dict[str, Any]]] = {}
        for item in clean_feed:
            c_id = item.get("cluster_id", "genel_gundem")
            clusters.setdefault(c_id, []).append(item)
            
        agenda_list = []
        for c_id, items in clusters.items():
            score = cls.calculate_cross_perspective_score(items)
            focus_item = items[0]
            
            agenda_list.append({
                "cluster_id": c_id,
                "focus_source": focus_item.get("source_name", "Bilinmeyen Kaynak"),
                "focus_news": focus_item.get("title", ""),
                "item_count": len(items),
                "kurdish_media_count": sum(1 for i in items if i.get("source_name") in KURDISH_MEDIA_SOURCES),
                "cross_perspective_score": score,
                "items": items
            })
            
        # Puanına göre sırala
        agenda_list.sort(key=lambda x: x["cross_perspective_score"], reverse=True)
        return agenda_list


# ==============================================================================
# 4. BÖLÜM: GEPHI NETWORK GRAPH EXPORTER (.gexf)
# ==============================================================================

class GephiExporter:
    """
    Analiz Sepetindeki verilerden Aktör, Kaynak ve Söylem ilişkisini içeren 
    .gexf formatında Gephi ağ dosyası üretir.
    """
    @staticmethod
    def export_gexf(analiz_sepeti_items: List[Dict[str, Any]], output_filepath: str = "analiz_sepeti.gexf") -> str:
        G = nx.DiGraph()

        for idx, item in enumerate(analiz_sepeti_items, 1):
            source = item.get("source_name", f"Kaynak_{idx}")
            discourse_group = item.get("discourse_segment", "Genel / Siyasi Çevre")
            category = item.get("category", "Siyasi Süreç / Diyalog")
            title = item.get("title", "Gelişme")

            # Düğümleri ekle
            G.add_node(source, label=source, node_type="Kaynak")
            G.add_node(discourse_group, label=discourse_group, node_type="Söylem Çevresi")

            # Kenar (ilişki) ekle
            G.add_edge(source, discourse_group, label=category, title=title, weight=1.0)

        nx.write_gexf(G, output_filepath, encoding='utf-8', prettyprint=True)
        logger.info(f"Gephi grafik dosyası başarıyla oluşturuldu: {output_filepath}")
        return output_filepath


# ==============================================================================
# 5. BÖLÜM: LLM RAPOR ÜRETİM PROMPTU (PDF Formatında Düz Yazı & Kesim Ayrımı)
# ==============================================================================

LLM_PDF_STYLE_SYSTEM_PROMPT = """
Sen üst düzey bir stratejik istihbarat ve güvenlik analistisin. 
Sana verilen 'Analiz Sepeti' haberlerini değerlendirerek tıpkı aşağıdaki kurallara sahip profesyonel bir PDF analiz belgesi gibi metin kaleme alacaksın.

KESİN FORMAT KURALLARI:
1. ŞABLON CÜMLE YASAĞI:
   "X kaynağında Y aktarılmaktadır", "Bu içerikte bilgilendirici yaklaşım öne çıkmaktadır" gibi mekanik kalıpları KESİNLİKLE KULLANMA.
2. DÜZ YAZI VE PARAGRAF DÜZENİ:
   - Metin akıcı, diplomatik ve akademik bir dille yazılmalıdır.
   - Her bir gelişme/haber TEK BİR PARAGRAF olarak ele alınmalıdır.
3. KESİM AYRIMLARI VE ATIFLAR (EN ÖNEMLİ UNSUR):
   - Paragraf başında olay/gelişme özetlenmeli, ardından "hangi kesim ne diyor?" ayrımı metnin içerisine doğal bir şekilde yedirilmelidir.
   - Kesim Ayrımları: 
     * Abdullah Öcalan'a biat eden kesim
     * Milliyetçi Kürt kesim / Barzani çizgisi (KDP/Rudaw)
     * Sol-Liberal / Aşırı Sol kesim (HBDH/MLKP vb.)
     * Devlet / Hükümet / Milliyetçi Türk kanadı
   - Metin içindeki iddia, tespit veya alıntıların sonuna ilgili kaynağın dipnot numarasını köşeli parantezle ekle (Örn: "...olduğu iddia edilmiştir[1].", "...ihanet olarak nitelendirilmiştir[3][4].").
4. KAYNAKÇA LİSTESİ:
   - Metnin en altında 1'den başlayarak tüm kaynakların No, Kaynak İsmi, Başlık ve URL bilgilerini içeren bir Kaynakça listesi oluştur.
"""

class LLMReportGenerator:
    """
    LLM API entegrasyonu sağlayan rapor üretici sınıf.
    """
    def __init__(self, llm_client_instance):
        self.client = llm_client_instance

    def generate_report(self, analiz_sepeti_items: List[Dict[str, Any]]) -> str:
        payload = []
        for idx, item in enumerate(analiz_sepeti_items, 1):
            payload.append({
                "no": idx,
                "source": item.get("source_name"),
                "title": item.get("title"),
                "summary": item.get("summary") or item.get("content"),
                "url": item.get("url")
            })

        user_content = f"Analiz Sepetindeki Haberler:\n{json.dumps(payload, ensure_ascii=False, indent=2)}"

        # LLM API Çağrısı
        response = self.client.generate_content(
            contents=[
                {"role": "user", "parts": [LLM_PDF_STYLE_SYSTEM_PROMPT, user_content]}
            ]
        )
        return response.text


# ==============================================================================
# 6. BÖLÜM: WORD (.DOCX) DOKÜMANI OLUŞTURUCUSU
# ==============================================================================

class WordDocumentBuilder:
    """
    LLM çıktısını hedeflenilen Word belgesine dönüştüren modül.
    """
    @staticmethod
    def build_docx(report_text: str, output_filepath: str = "Terorsuz_Turkiye_Analiz_Raporu.docx") -> str:
        doc = Document()

        # Sayfa Kenar Boşlukları
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Başlık
        p_head = doc.add_paragraph()
        p_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_head = p_head.add_run("TERÖRSÜZ TÜRKİYE - AÇIK KAYNAK ANALİZ RAPORU\n")
        r_head.bold = True
        r_head.font.name = "Arial"
        r_head.font.size = Pt(15)
        r_head.font.color.rgb = RGBColor(0, 32, 96)

        # Tarih
        p_date = doc.add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_date = p_date.add_run(datetime.now().strftime("%d.%m.%Y") + "\n")
        r_date.font.name = "Arial"
        r_date.font.size = Pt(10)
        r_date.font.italic = True
        r_date.font.color.rgb = RGBColor(89, 89, 89)

        # Metin Paragrafları
        blocks = report_text.split("\n\n")
        for block in blocks:
            text = block.strip()
            if not text:
                continue

            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(8)

            # Kaynakça başlığı kontrolü
            if text.startswith("Kaynakça") or text.startswith("Kaynaklar"):
                run = p.add_run(text)
                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 32, 96)
            else:
                run = p.add_run(text)
                run.font.name = "Arial"
                run.font.size = Pt(10.5)

        doc.save(output_filepath)
        logger.info(f"Word raporu kaydedildi: {output_filepath}")
        return output_filepath


# ==============================================================================
# 7. BÖLÜM: TÜM SİSTEMİ ÇALIŞTIRAN ANA PIPELINE
# ==============================================================================

class TerorsuzTurkiyePipeline:
    def __init__(self, llm_client):
        self.llm_client = llm_client
        self.report_gen = LLMReportGenerator(llm_client)

    def run(self, raw_delta_feed: List[Dict[str, Any]], 
            all_scraped_feed: List[Dict[str, Any]], 
            selected_analiz_sepeti: List[Dict[str, Any]]) -> Dict[str, Any]:
        
        logger.info("Pipeline başlatıldı...")

        # 1. Delta Tarama Filtreleme (Scary Movie vb. Temizliği)
        clean_delta = ContentFilter.filter_delta_scan(raw_delta_feed)

        # 2. Gündem ve Söylem Haritası Oluşturma
        agenda_map = DiscourseClusteringEngine.build_agenda_map(all_scraped_feed)

        # 3. PDF Stilinde Analiz Sepeti Rapor Metni Üretimi
        report_text = self.report_gen.generate_report(selected_analiz_sepeti)

        # 4. Word Belgesi Oluşturma
        docx_path = WordDocumentBuilder.build_docx(report_text, "Terorsuz_Turkiye_Analiz_Raporu.docx")

        # 5. Gephi Grafiği Oluşturma
        gephi_path = GephiExporter.export_gexf(selected_analiz_sepeti, "analiz_sepeti.gexf")

        logger.info("Pipeline başarıyla tamamlandı.")

        return {
            "clean_delta_items": clean_delta,
            "agenda_map": agenda_map,
            "report_text": report_text,
            "docx_path": docx_path,
            "gephi_path": gephi_path
        }
